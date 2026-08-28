import os
import sys
import re

# התאמת קידוד פלט ל-UTF-8 בסביבת Windows
if sys.stdout:
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

HOUR_MAPPING = {
    '1': 'שעה 1 (08:00-08:50)',
    '2': 'שעה 2 (08:50-09:35)',
    '3': 'שעה 3 (10:15-11:00)',
    '4': 'שעה 4 (11:00-11:45)',
    '5': 'שעה 5 (12:00-12:45)',
    '6': 'שעה 6 (12:45-13:30)',
    '7': 'שעה 7 (13:30-14:20)',
    '8': 'שעה 8',
    '9': 'שעה 9',
}

DAYS_ORDER = [
    ('ראשון', 'יום ראשון'),
    ('שני', 'יום שני'),
    ('שלישי', 'יום שלישי'),
    ('רביעי', 'יום רביעי'),
    ('חמישי', 'יום חמישי'),
    ('שישי', 'יום שישי'),
]


def set_cell_margins(cell, top=5, bottom=5, left=50, right=50):
    """הגדרת שוליים לתא בטבלה"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    """הגדרת צבע רקע לתא"""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shd)


def set_rtl(paragraph_or_cell):
    """הגדרת כיווניות RTL"""
    pPr = paragraph_or_cell._element.get_or_add_pPr() if hasattr(paragraph_or_cell, '_element') else paragraph_or_cell
    if pPr.find(qn('w:bidi')) is None:
        pPr.append(OxmlElement('w:bidi'))


def normalize_class_name(name):
    """נרמול שם כיתה להשוואה (למשל: כיתה א'1 -> א1)"""
    return name.replace('מערכת שעות', '').replace('כיתה', '').replace("'", '').replace('"', '').replace(' ', '').strip()


def parse_teachers_file(teachers_docx_path):
    """
    חילוץ כל השיעורים המלאים מקובץ מורים טבלה:
    החזרה:
      teacher_schedules: רשימת מערכות שעות של מורים
      class_lessons_map: מיפוי (class_norm, day_key, hour_str) -> רשימת שיעורים מלאים (מקצוע מלא, שם מורה מלא, סוג)
    """
    doc_t = Document(teachers_docx_path)
    teacher_schedules = []
    class_lessons_map = {}

    elements = []
    for el in doc_t.element.body:
        if el.tag.endswith('tbl'):
            from docx.table import Table
            elements.append(('tbl', Table(el, doc_t)))
        elif el.tag.endswith('p'):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, doc_t)
            if p.text.strip():
                elements.append(('p', p.text.strip()))

    i = 0
    while i < len(elements):
        if (i + 2 < len(elements) and 
            elements[i][0] == 'tbl' and 
            elements[i+1][0] == 'p' and 
            elements[i+2][0] == 'tbl'):
            
            header_tbl = elements[i][1]
            title_text = elements[i+1][1]
            teacher_name = title_text.replace('מערכת שעות למורה', '').replace('מערכת שעות מורה', '').strip()
            grid_tbl = elements[i+2][1]
            
            header_row = [c.text.strip() for c in grid_tbl.rows[0].cells]
            col_indices = {}
            for c_idx, col_name in enumerate(header_row):
                for day_key, _ in DAYS_ORDER:
                    if day_key in col_name:
                        col_indices[day_key] = c_idx
            hour_col_idx = 6 if len(header_row) > 6 else len(header_row) - 1
            
            teacher_day_schedule = {day_key: [] for day_key, _ in DAYS_ORDER}

            for row in grid_tbl.rows[1:]:
                hour_val = row.cells[hour_col_idx].text.strip()
                hour_str = HOUR_MAPPING.get(hour_val, f'שעה {hour_val}')
                
                for day_key, _ in DAYS_ORDER:
                    if day_key in col_indices:
                        col_i = col_indices[day_key]
                        raw_text = row.cells[col_i].text.strip()
                        if raw_text:
                            lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
                            subj = lines[0] if len(lines) > 0 else ''
                            cls = lines[1] if len(lines) > 1 else ''
                            act = lines[2] if len(lines) > 2 else ''
                            
                            formatted_lesson = ', '.join(lines)
                            teacher_day_schedule[day_key].append((hour_str, formatted_lesson))
                            
                            # אם זה שיעור של כיתה (למשל ה'1, א'3...)
                            if cls and any(c in cls for c in ['א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י']):
                                norm_c = normalize_class_name(cls)
                                key = (norm_c, day_key, hour_val)
                                lesson_tuple = (subj, teacher_name, act)
                                if lesson_tuple not in class_lessons_map.setdefault(key, []):
                                    class_lessons_map[key].append(lesson_tuple)

            date_str = header_tbl.rows[0].cells[0].text.strip() if len(header_tbl.rows[0].cells) > 0 else ''
            school_str = header_tbl.rows[0].cells[1].text.strip() if len(header_tbl.rows[0].cells) > 1 else 'יצחק נבון'
            teacher_schedules.append({
                'date_str': date_str,
                'school_str': school_str,
                'title_text': title_text,
                'schedule': teacher_day_schedule
            })

            i += 3
        else:
            i += 1

    return teacher_schedules, class_lessons_map


def build_list_document(output_docx_path, schedules_data):
    """בונה מסמך DOCX בפורמט רשימות (סגנון קורצאק)"""
    out_doc = Document()

    for section in out_doc.sections:
        section.top_margin = Pt(11)
        section.bottom_margin = Pt(22)
        section.left_margin = Pt(22)
        section.right_margin = Pt(22)

    for idx, item in enumerate(schedules_data):
        # 1. טבלת מטא (תאריך ושם בית ספר)
        meta_table = out_doc.add_table(rows=1, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        c0 = meta_table.rows[0].cells[0]
        c0.text = item.get('date_str', '')
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.runs[0].font.name = 'Arial'
        p0.runs[0].font.size = Pt(10)
        
        c1 = meta_table.rows[0].cells[1]
        c1.text = item.get('school_str', 'יצחק נבון')
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.runs[0].font.name = 'Arial'
        p1.runs[0].font.size = Pt(10)

        # 2. כותרת המערכת
        title_p = out_doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = title_p.add_run(item.get('title_text', ''))
        t_run.font.name = 'Arial'
        t_run.font.size = Pt(14)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(0, 0, 0)
        set_rtl(title_p)

        # 3. טבלת הרשימה
        day_schedule = item.get('schedule', {})
        total_rows = 0
        active_days = []
        for day_key, day_title in DAYS_ORDER:
            lessons = day_schedule.get(day_key, [])
            if lessons:
                active_days.append((day_title, lessons))
                total_rows += 1 + len(lessons)

        if total_rows > 0:
            list_table = out_doc.add_table(rows=total_rows, cols=2)
            list_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}><w:gridCol w:w="6000"/><w:gridCol w:w="2000"/></w:tblGrid>')
            list_table._element.insert(1, tblGrid)

            curr_row_idx = 0
            for day_title, lessons in active_days:
                # שורת יום
                day_row = list_table.rows[curr_row_idx]
                c_left = day_row.cells[0]
                set_cell_shading(c_left, '28486B')
                set_cell_margins(c_left)
                set_rtl(c_left.paragraphs[0])
                
                c_right = day_row.cells[1]
                set_cell_shading(c_right, '28486B')
                set_cell_margins(c_right)
                p_r = c_right.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r_day = p_r.add_run(day_title)
                r_day.font.name = 'Arial'
                r_day.font.size = Pt(9)
                r_day.font.bold = True
                r_day.font.color.rgb = RGBColor(255, 255, 255)
                set_rtl(p_r)
                
                curr_row_idx += 1

                # שורות שיעורים
                for hour_label, lesson_content in lessons:
                    lesson_row = list_table.rows[curr_row_idx]
                    
                    c_lesson = lesson_row.cells[0]
                    set_cell_shading(c_lesson, 'FFFFFF')
                    set_cell_margins(c_lesson)
                    p_les = c_lesson.paragraphs[0]
                    p_les.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_les = p_les.add_run(lesson_content)
                    r_les.font.name = 'Arial'
                    r_les.font.size = Pt(9)
                    r_les.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_les)

                    c_hour = lesson_row.cells[1]
                    set_cell_shading(c_hour, 'FFFFFF')
                    set_cell_margins(c_hour)
                    p_hr = c_hour.paragraphs[0]
                    p_hr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_hr = p_hr.add_run(hour_label)
                    r_hr.font.name = 'Arial'
                    r_hr.font.size = Pt(9)
                    r_hr.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_hr)

                    curr_row_idx += 1

        if idx < len(schedules_data) - 1:
            p_break = out_doc.add_paragraph()
            p_break.add_run().add_break(WD_BREAK.PAGE)

    out_doc.save(output_docx_path)
    print(f"נוצר בהצלחה: {output_docx_path}")


def parse_classes_file(classes_docx_path, class_lessons_map):
    """
    חילוץ מערכות שעות של כיתות, תוך העשרה והשלמה של שמות מלאים מקובץ המורים.
    במידה ויש פיצול מקבילי (כמו אומנות ירוקה בגינה + תיאטרון), כל שיעור נרשם בשורה נפרדת בתא.
    """
    doc_c = Document(classes_docx_path)
    class_schedules = []

    elements = []
    for el in doc_c.element.body:
        if el.tag.endswith('tbl'):
            from docx.table import Table
            elements.append(('tbl', Table(el, doc_c)))
        elif el.tag.endswith('p'):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, doc_c)
            if p.text.strip():
                elements.append(('p', p.text.strip()))

    i = 0
    while i < len(elements):
        if (i + 2 < len(elements) and 
            elements[i][0] == 'tbl' and 
            elements[i+1][0] == 'p' and 
            elements[i+2][0] == 'tbl'):
            
            header_tbl = elements[i][1]
            title_text = elements[i+1][1]
            norm_class = normalize_class_name(title_text)
            grid_tbl = elements[i+2][1]
            
            header_row = [c.text.strip() for c in grid_tbl.rows[0].cells]
            col_indices = {}
            for c_idx, col_name in enumerate(header_row):
                for day_key, _ in DAYS_ORDER:
                    if day_key in col_name:
                        col_indices[day_key] = c_idx
            hour_col_idx = 6 if len(header_row) > 6 else len(header_row) - 1
            
            class_day_schedule = {day_key: [] for day_key, _ in DAYS_ORDER}

            for row in grid_tbl.rows[1:]:
                hour_val = row.cells[hour_col_idx].text.strip()
                hour_str = HOUR_MAPPING.get(hour_val, f'שעה {hour_val}')
                
                for day_key, _ in DAYS_ORDER:
                    if day_key in col_indices:
                        col_i = col_indices[day_key]
                        raw_text = row.cells[col_i].text.strip()
                        if raw_text:
                            # בדיקה אם יש לנו שיעור מלא מקובץ מורים
                            key = (norm_class, day_key, hour_val)
                            teacher_lessons = class_lessons_map.get(key, [])
                            
                            if teacher_lessons:
                                # בניית שורות מלאות ומדויקות לכל מורה / מקצוע
                                lesson_lines = []
                                for subj, teacher, act in teacher_lessons:
                                    parts = [p for p in [subj, teacher, act] if p]
                                    lesson_lines.append(', '.join(parts))
                                final_lesson_text = '\n'.join(lesson_lines)
                            else:
                                # ברירת מחדל מהטבלה אם מורה חיצוני/צהרון
                                lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
                                final_lesson_text = ', '.join(lines)
                            
                            class_day_schedule[day_key].append((hour_str, final_lesson_text))

            date_str = header_tbl.rows[0].cells[0].text.strip() if len(header_tbl.rows[0].cells) > 0 else ''
            school_str = header_tbl.rows[0].cells[1].text.strip() if len(header_tbl.rows[0].cells) > 1 else 'יצחק נבון'
            class_schedules.append({
                'date_str': date_str,
                'school_str': school_str,
                'title_text': title_text,
                'schedule': class_day_schedule
            })

            i += 3
        else:
            i += 1

    return class_schedules


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    teachers_input = os.path.join(base_dir, "מורים טבלה.docx")
    teachers_output = os.path.join(base_dir, "מורים.docx")
    classes_input = os.path.join(base_dir, "כיתות טבלה.docx")
    classes_output = os.path.join(base_dir, "כיתות.docx")
    
    if os.path.exists(teachers_input):
        print("מעבד קובץ מורים טבלה ומחלץ שמות מלאים של כל המקצועות והמורים...")
        teacher_schedules, class_lessons_map = parse_teachers_file(teachers_input)
        build_list_document(teachers_output, teacher_schedules)
    else:
        teacher_schedules, class_lessons_map = [], {}

    if os.path.exists(classes_input):
        print("מעבד קובץ כיתות טבלה ומסנכרן שמות מלאים מקובץ מורים...")
        class_schedules = parse_classes_file(classes_input, class_lessons_map)
        build_list_document(classes_output, class_schedules)


if __name__ == "__main__":
    main()
