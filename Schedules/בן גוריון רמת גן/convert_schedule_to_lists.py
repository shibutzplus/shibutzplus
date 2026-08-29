import os
import sys
import re
from datetime import datetime

# התאמת קידוד פלט ל-UTF-8 בסביבת Windows
if sys.stdout:
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

HOUR_MAPPING = {
    '1': 'שעה 1 (08:00-08:50)',
    '2': 'שעה 2 (08:50-09:35)',
    '3': 'שעה 3 (10:15-11:00)',
    '4': 'שעה 4 (11:00-11:45)',
    '5': 'שעה 5 (12:00-12:45)',
    '6': 'שעה 6 (12:45-13:30)',
    '7': 'שעה 7 (13:30-14:20)',
    '8': 'שעה 8',
    '9': 'שעה 9',
}

DAYS_ORDER = [
    ('ראשון', 'יום ראשון'),
    ('שני', 'יום שני'),
    ('שלישי', 'יום שלישי'),
    ('רביעי', 'יום רביעי'),
    ('חמישי', 'יום חמישי'),
    ('שישי', 'יום שישי'),
]

ORDERED_CLASS_CODES = [
    'א1', 'א2', 'א3',
    'ב1', 'ב2',
    'ג1', 'ג2',
    'ד1', 'ד2', 'ד3',
    'ה1', 'ה2',
    'ו1', 'ו2',
    'ז1', 'ז2', 'ז3',
    'ח1', 'ח2', 'ח3', 'ח4'
]

# מיפוי מקצועות, קיצורים ושמות נרדפים לשם המקצוע התקני
SUBJECT_CANONICAL_MAP = {
    'ניצנים של בינה': 'ניצנים של בינה',
    'חינוך פיננסי': 'חינוך פיננסי',
    'אנגלית מדוברת': 'אנגלית מדוברת',
    'קריאה מונחית': 'קריאה מונחית',
    'כישורי חיים': 'כישורי חיים',
    'כלי מיתר': 'כלי מיתר',
    'רובוטיקה': 'רובוטיקה',
    'מתמטיקה': 'מתמטיקה',
    'מתמתיקה': 'מתמטיקה',
    'ספרות': 'ספרות',
    'מדעים': 'מדעים',
    'אנגלית': 'אנגלית',
    'ערבית': 'ערבית',
    'ספורט': 'ספורט',
    'שחמט': 'שחמט',
    'מגמות': 'מגמות',
    'דרמה': 'דרמה',
    'אשכול': 'אשכול',
    'תנ"ך': 'תנ"ך',
    'תנך': 'תנ"ך',
}

# רשימת המקצועות לחיפוש (ממוינת מארוך לקצר)
KNOWN_SUBJECTS = sorted(SUBJECT_CANONICAL_MAP.keys(), key=len, reverse=True)

# תוכניות עצמאיות (כאשר מופיעות לבד בתא, שם המורה הוא שם המקצוע)
STANDALONE_PROGRAMS = [
    'אנגלית מדוברת',
    'רובוטיקה',
    'שחמט',
    'כלי מיתר',
    'מגמות',
]


def set_cell_margins(cell, top=5, bottom=5, left=50, right=50):
    """הגדרת שוליים לתא בטבלה"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    """הגדרת צבע רקע לתא"""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shd)


def set_rtl(paragraph_or_cell):
    """הגדרת כיווניות RTL"""
    pPr = paragraph_or_cell._element.get_or_add_pPr() if hasattr(paragraph_or_cell, '_element') else paragraph_or_cell
    if pPr.find(qn('w:bidi')) is None:
        pPr.append(OxmlElement('w:bidi'))


def parse_single_entry(entry_text):
    """
    מנתח קטע טקסט בודד ומחלץ מקצוע, מורה וסוג פעילות
    """
    clean = entry_text.strip(' ,-')
    if not clean:
        return []
    
    # 1. בדיקת פרטני
    if 'פרטני' in clean:
        teacher = clean.replace('פרטני', '').replace(':', '').replace('-', '').strip()
        teacher = teacher if teacher else 'פרטני'
        return [{
            'subject': 'פרטני',
            'teacher': teacher,
            'activity': 'פרטני'
        }]
    
    # 2. בדיקת תוכניות עצמאיות
    for sp in STANDALONE_PROGRAMS:
        if clean == sp:
            return [{
                'subject': sp,
                'teacher': sp,
                'activity': 'הוראה'
            }]
            
    # 3. בדיקת חינוך פיננסי / עננים
    if 'חינוך פיננסי' in clean or 'עננים' in clean:
        return [{
            'subject': 'חינוך פיננסי',
            'teacher': 'עננים',
            'activity': 'הוראה'
        }]
        
    # 4. חיפוש מקצוע מתוך הרשימה המוגדרת מראש
    for subj in KNOWN_SUBJECTS:
        pattern = rf'(?:[\s\-_]+|^){re.escape(subj)}(?:[\s\-_]+|$)'
        m = re.search(pattern, clean)
        if m:
            teacher_part = (clean[:m.start()] + " " + clean[m.end():]).strip(' -_:,')
            if not teacher_part:
                teacher_part = subj
                
            final_subj = SUBJECT_CANONICAL_MAP.get(subj, subj)
            if teacher_part == 'נוע לי':
                teacher_part = 'נטע לי'
                
            return [{
                'subject': final_subj,
                'teacher': teacher_part,
                'activity': 'הוראה'
            }]
            
    # 5. ללא מקצוע מזוהה - המקצוע "-" והמורה הוא הטקסט
    teacher_clean = clean.strip(' -_:,')
    if teacher_clean == 'נוע לי':
        teacher_clean = 'נטע לי'
        
    return [{
        'subject': '-',
        'teacher': teacher_clean,
        'activity': 'הוראה'
    }]


def parse_cell_lessons(raw_text):
    """
    פירוק תא שיעור במערכת לרשימת שיבוצים
    """
    if not raw_text or not raw_text.strip():
        return []
    
    text = raw_text.replace('/', '\n').replace(';', '\n')
    raw_lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    lessons = []
    for line in raw_lines:
        # בדיקה אם יש פיצול של מספר מורים באמצעות '+'
        if '+' in line:
            parts = [p.strip() for p in line.split('+') if p.strip()]
            for part in parts:
                lessons.extend(parse_single_entry(part))
        # פיצול זוגי של חווה-אורן או חווה-רון
        elif line.replace(' ', '') in ['חווה-אורן', 'חווה-רון']:
            parts = [p.strip() for p in line.split('-') if p.strip()]
            for part in parts:
                lessons.extend(parse_single_entry(part))
        else:
            lessons.extend(parse_single_entry(line))
            
    return lessons


def parse_schedule_file(input_docx_path):
    """
    קריאת קובץ מערכת השעות המקורית וחילוץ נתוני הכיתות והמורים
    """
    doc = Document(input_docx_path)
    
    classes_data = []
    current_paragraphs = []
    table_index = 0
    
    for el in doc.element.body:
        if el.tag.endswith('p'):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, doc)
            t = p.text.strip()
            if t:
                current_paragraphs.append(t)
        elif el.tag.endswith('tbl'):
            from docx.table import Table
            tbl = Table(el, doc)
            
            # בדיקה האם לטבלה יש נתונים (ולא טבלה ריקה)
            has_data = any(any(c.text.strip() for c in r.cells[1:]) for r in tbl.rows[1:])
            if not has_data:
                current_paragraphs = []
                continue
            
            # קביעת שם הכיתה לפי המיפוי המסודר
            if table_index < len(ORDERED_CLASS_CODES):
                class_code = ORDERED_CLASS_CODES[table_index]
            else:
                raw_title = " ".join(current_paragraphs)
                m = re.search(r'([א-ח][׳\']?\s*[1-9])', raw_title)
                class_code = m.group(1).replace("'", "").replace(" ", "") if m else f"כיתה_{table_index+1}"
            
            class_name = f"כיתה {class_code}"
            class_title = f"מערכת שעות {class_name}"
            
            header_row = [c.text.strip() for c in tbl.rows[0].cells]
            day_col_map = {}
            for c_idx, val in enumerate(header_row):
                for day_key, _ in DAYS_ORDER:
                    if day_key in val:
                        day_col_map[day_key] = c_idx
                        
            day_schedule = {day_key: [] for day_key, _ in DAYS_ORDER}
            
            for row in tbl.rows[1:]:
                hour_val = row.cells[0].text.strip()
                if not hour_val:
                    continue
                hour_str = HOUR_MAPPING.get(hour_val, f"שעה {hour_val}")
                
                for day_key, _ in DAYS_ORDER:
                    if day_key in day_col_map:
                        col_idx = day_col_map[day_key]
                        cell_txt = row.cells[col_idx].text.strip()
                        if cell_txt:
                            lessons = parse_cell_lessons(cell_txt)
                            if lessons:
                                day_schedule[day_key].append((hour_val, hour_str, lessons))
                                
            classes_data.append({
                'class_code': class_code,
                'class_name': class_name,
                'title_text': class_title,
                'schedule': day_schedule
            })
            
            current_paragraphs = []
            table_index += 1
            
    return classes_data


def build_teacher_schedules(classes_data):
    """
    בניית מערכות שעות לכל המורים על בסיס הצלבת מערכות השעות של הכיתות
    """
    teachers_map = {}
    
    for cls_item in classes_data:
        class_name = cls_item['class_name']
        day_schedule = cls_item['schedule']
        
        for day_key, _ in DAYS_ORDER:
            lessons_in_day = day_schedule.get(day_key, [])
            for hour_val, hour_str, lessons in lessons_in_day:
                for lesson in lessons:
                    teacher = lesson['teacher']
                    subject = lesson['subject']
                    act = lesson['activity']
                    
                    if not teacher or teacher == 'פרטני':
                        continue
                        
                    if teacher not in teachers_map:
                        teachers_map[teacher] = {day_k: {} for day_k, _ in DAYS_ORDER}
                        
                    if hour_str not in teachers_map[teacher][day_key]:
                        teachers_map[teacher][day_key][hour_str] = []
                        
                    lesson_desc = f"{subject}, {class_name}, {act}"
                    if lesson_desc not in teachers_map[teacher][day_key][hour_str]:
                        teachers_map[teacher][day_key][hour_str].append(lesson_desc)
                        
    # המרה למבנה מסודר לרשימת מורים
    teachers_data = []
    for teacher_name in sorted(teachers_map.keys()):
        teacher_day_schedule = {day_k: [] for day_k, _ in DAYS_ORDER}
        for day_key, _ in DAYS_ORDER:
            hours_dict = teachers_map[teacher_name][day_key]
            for hour_val in sorted(HOUR_MAPPING.keys(), key=lambda x: int(x)):
                hour_str = HOUR_MAPPING[hour_val]
                if hour_str in hours_dict:
                    lessons_list = hours_dict[hour_str]
                    combined_text = "\n".join(lessons_list)
                    teacher_day_schedule[day_key].append((hour_str, combined_text))
                    
        teachers_data.append({
            'teacher_name': teacher_name,
            'title_text': f"מערכת שעות מורה {teacher_name}",
            'schedule': teacher_day_schedule
        })
        
    return teachers_data


def generate_docx_document(output_path, items_data, school_name="בן גוריון", is_classes=True):
    """
    יצירת קובץ DOCX מעוצב בפורמט קורצאק המדויק
    """
    doc = Document()
    
    # הגדרת שולי עמוד
    for section in doc.sections:
        section.top_margin = Pt(11)
        section.bottom_margin = Pt(22)
        section.left_margin = Pt(22)
        section.right_margin = Pt(22)
        
    now_str = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    
    for idx, item in enumerate(items_data):
        # 1. טבלת מטא (תאריך ושם בית ספר)
        meta_tbl = doc.add_table(rows=1, cols=2)
        meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        c0 = meta_tbl.rows[0].cells[0]
        c0.text = now_str
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.runs[0].font.name = 'Arial'
        p0.runs[0].font.size = Pt(10)
        
        c1 = meta_tbl.rows[0].cells[1]
        c1.text = school_name
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.runs[0].font.name = 'Arial'
        p1.runs[0].font.size = Pt(10)
        
        # 2. כותרת המערכת
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = title_p.add_run(item['title_text'])
        t_run.font.name = 'Arial'
        t_run.font.size = Pt(14)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(0, 0, 0)
        set_rtl(title_p)
        
        # 3. טבלת רשימת שיעורים
        day_schedule = item['schedule']
        total_rows = 0
        active_days = []
        
        for day_key, day_title in DAYS_ORDER:
            if is_classes:
                raw_lessons = day_schedule.get(day_key, [])
                formatted_lessons = []
                for hour_val, hour_str, lessons in raw_lessons:
                    lines = [f"{l['subject']}, {l['teacher']}, {l['activity']}" for l in lessons]
                    formatted_lessons.append((hour_str, "\n".join(lines)))
            else:
                formatted_lessons = day_schedule.get(day_key, [])
                
            if formatted_lessons:
                active_days.append((day_title, formatted_lessons))
                total_rows += 1 + len(formatted_lessons)
                
        if total_rows > 0:
            list_table = doc.add_table(rows=total_rows, cols=2)
            list_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}><w:gridCol w:w="6000"/><w:gridCol w:w="2000"/></w:tblGrid>')
            list_table._element.insert(1, tblGrid)
            
            curr_row_idx = 0
            for day_title, lessons in active_days:
                # שורת כותרת יום (#28486B)
                day_row = list_table.rows[curr_row_idx]
                c_left = day_row.cells[0]
                set_cell_shading(c_left, '28486B')
                set_cell_margins(c_left)
                set_rtl(c_left.paragraphs[0])
                
                c_right = day_row.cells[1]
                set_cell_shading(c_right, '28486B')
                set_cell_margins(c_right)
                p_r = c_right.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r_day = p_r.add_run(day_title)
                r_day.font.name = 'Arial'
                r_day.font.size = Pt(9)
                r_day.font.bold = True
                r_day.font.color.rgb = RGBColor(255, 255, 255)
                set_rtl(p_r)
                
                curr_row_idx += 1
                
                # שורות השיעורים
                for hour_label, lesson_content in lessons:
                    lesson_row = list_table.rows[curr_row_idx]
                    
                    # תא שמאל - תוכן השיעור
                    c_lesson = lesson_row.cells[0]
                    set_cell_shading(c_lesson, 'FFFFFF')
                    set_cell_margins(c_lesson)
                    p_les = c_lesson.paragraphs[0]
                    p_les.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_les = p_les.add_run(lesson_content)
                    r_les.font.name = 'Arial'
                    r_les.font.size = Pt(9)
                    r_les.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_les)
                    
                    # תא ימין - תווית השעה
                    c_hour = lesson_row.cells[1]
                    set_cell_shading(c_hour, 'FFFFFF')
                    set_cell_margins(c_hour)
                    p_hr = c_hour.paragraphs[0]
                    p_hr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_hr = p_hr.add_run(hour_label)
                    r_hr.font.name = 'Arial'
                    r_hr.font.size = Pt(9)
                    r_hr.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_hr)
                    
                    curr_row_idx += 1
                    
        # מעבר עמוד בין כיתות/מורים
        if idx < len(items_data) - 1:
            p_break = doc.add_paragraph()
            p_break.add_run().add_break(WD_BREAK.PAGE)
            
    doc.save(output_path)
    print(f"נוצר בהצלחה: {output_path}")


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_file = os.path.join(base_dir, "מערכת שעות.docx")
    
    classes_output = os.path.join(base_dir, "כיתות.docx")
    teachers_output = os.path.join(base_dir, "מורים.docx")
    
    if not os.path.exists(input_file):
        print(f"שגיאה: הקובץ {input_file} אינו קיים.")
        return
        
    print("מעבד את קובץ מערכת השעות של בית ספר בן גוריון...")
    classes_data = parse_schedule_file(input_file)
    print(f"חולצו {len(classes_data)} כיתות.")
    
    print("\nיוצר קובץ כיתות.docx...")
    generate_docx_document(classes_output, classes_data, school_name="בן גוריון", is_classes=True)
    
    print("\nיוצר קובץ מורים.docx...")
    teachers_data = build_teacher_schedules(classes_data)
    print(f"חולצו {len(teachers_data)} מורים.")
    generate_docx_document(teachers_output, teachers_data, school_name="בן גוריון", is_classes=False)
    
    print("\nההמרה הושלמה בהצלחה!")


if __name__ == "__main__":
    main()
