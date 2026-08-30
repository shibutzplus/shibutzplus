"""
Converter for "Sharret Kiryat Ono" School Schedules
School ID: nmjjwezgwhmlwfk3p10zwztf

Converts:
  1) מערכת מקור כיתות.pdf -> כיתות.docx
  2) מערכת מקור מורים.pdf -> מורים.docx

Leverages database reference entities (Teachers, Classes, Subjects, WorkGroups)
for exact, canonical matching across PDF text, reversed strings, and variations.
"""

import os
import sys
import re
import pymupdf
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CLASS_PDF = os.path.join(BASE_DIR, "מערכת מקור כיתות.pdf")
TEACHER_PDF = os.path.join(BASE_DIR, "מערכת מקור מורים.pdf")
CLASS_DOCX = os.path.join(BASE_DIR, "כיתות.docx")
TEACHER_DOCX = os.path.join(BASE_DIR, "מורים.docx")

# School metadata
SCHOOL_ID = "nmjjwezgwhmlwfk3p10zwztf"
SCHOOL_NAME = "משה שרת"
META_TIMESTAMP = "28/08/2026 13:38:55"

# Hour standard labels
HOUR_LABELS = {
    1: "שעה 1 (08:00-08:50)",
    2: "שעה 2 (08:50-09:35)",
    3: "שעה 3 (10:15-11:00)",
    4: "שעה 4 (11:00-11:45)",
    5: "שעה 5 (12:00-12:45)",
    6: "שעה 6 (12:45-13:30)",
    7: "שעה 7 (13:30-14:20)",
    8: "שעה 8",
    9: "שעה 9",
    10: "שעה 10",
    11: "שעה 11",
    12: "שעה 12"
}

DAYS_ORDER = [
    (1, "יום ראשון"),
    (2, "יום שני"),
    (3, "יום שלישי"),
    (4, "יום רביעי"),
    (5, "יום חמישי"),
    (6, "יום שישי")
]

# Database reference teachers for school nmjjwezgwhmlwfk3p10zwztf
DB_TEACHERS = [
    'אביגיל ברדה', 'אוראל עבאדי', 'אורית קסירר', 'אושר כהן', 'אושרה נדבורני',
    'איילת אדורם', 'איילת מלאך בוכניק', 'איילת מורד', 'אלקה הרשקוביץ', 'ארבל אלחרר',
    'אריאלה גרינברג', 'אתי יוסף', 'בית ספר מנגן - ג חליליות', 'גוני בן ברית',
    'גלית סמין', 'עינב דוד', 'דפני אלביני', 'הילה אלגן', 'חגית כהן',
    'חווה חקלאית ד', 'ימי חינוך ה', 'חן פניני', 'טלי נוראל', 'יערה מגן',
    'יפה דביר', 'יפה מזור', 'ישראל שוסטרמן', 'ליהיא שני סנדיק', 'לילך חייט',
    'מהנדסי הדור הבא', 'מורה כדור יד', 'מורה תכנות ו', 'מורה תל"ן הגנה עצמית',
    'מורה תל"ן הופ מדע א', 'מורה תל"ן כלבנות', 'תל"ן רופאים צעירים',
    'מיכל אומן', 'נאווה וייס', 'נילי שלפר', 'חן סגל', 'סמדר ברגר',
    'ספיר שירזי', 'עדי עבדה', 'עדי שקד', 'עמית ברק', 'מיה עמר',
    "שרה פרג'", 'חמי פרידמן', 'עמית פרל', 'רוני בועז', 'רותם איגמנוב',
    'רינה אשואל', 'ריקי תיאטרון', 'שני שם טוב', 'סיוון ורונסקי', 'מיה מולנר'
]

# Database reference classes and homerooms
CLASS_HOMEROOM = {
    '1א': 'חן סגל', '2א': 'עינב כהן', '3א': "שרה פרג'", '4א': 'ספיר שירזי',
    '1ב': 'אוראל עבאדי', '2ב': 'יערה מגן', '3ב': 'דפני אלביני', '4ב': 'עדי שקד',
    '1ג': 'מיה עמר', '2ג': 'עמית פרל', '3ג': 'נילי שלפר',
    '1ד': 'רוני בועז', '2ד': 'איילת מורד', '3ד': 'אלקה הרשקוביץ', '4ד': 'ספיר שירזי',
    '1ה': 'אריאלה גרינברג', '2ה': 'טלי נוראל', '3ה': 'חגית כהן', '4ה': 'טלי נוראל',
    '1ו': 'ישראל שוסטרמן', '2ו': 'איילת מלאך בוכניק', '3ו': 'הילה אלגן'
}

# DB-aligned Subject & WorkGroup corrections
PHRASE_CORRECTIONS = {
    'ג"חנ': 'חינוך גופני',
    'חנ"ג': 'חינוך גופני',
    'א"מתי': 'מתי"א',
    'מתי"א': 'מתי"א',
    'חיים כישורי': 'כישורי חיים',
    'הלב מפתח': 'מפתח הלב',
    'מורשת שבילי': 'שבילי מורשת',
    'ך"תנ': 'תנ"ך',
    'ספרייה העשרה': 'העשרה ספרייה',
    'סלים שעת': 'שעת סלים',
    'ב"זה': 'זה"ב',
    'שפה עברית': 'עברית',
    'מדעים וטכנולוגיה': 'מדעים',
    'מדע וטכנולוגיה': 'מדעים',
    'עצמית הגנה': 'הגנה עצמית',
    'הופ מדע': 'מדע הופ'
}

SPECIAL_TEACHER_TITLES = {
    'חליליות -ג מנגן ספר בית': 'בית ספר מנגן - ג חליליות',
    'ד חקלאית חווה': 'חווה חקלאית ד',
    'ה ימי חינוך': 'ימי חינוך ה',
    'הבא הדור מהנדסי': 'מהנדסי הדור הבא',
    'יד כדור מורה': 'מורה כדור יד',
    'ו תכנות מורה': 'מורה תכנות ו',
    'עצמית הגנה ן"תל מורה': 'מורה תל"ן הגנה עצמית',
    'א מדע הופ ן"תל מורה': 'מורה תל"ן הופ מדע א',
    'כלבנות ן"תל מורה': 'מורה תל"ן כלבנות',
    'תיאטרון ריקי': 'ריקי תיאטרון',
    'צעירים רופאים ן"תל': 'תל"ן רופאים צעירים',
    '3א שרה \'פרג': "שרה פרג'",
    'חנמ 1ב שני טוב שם': 'שני שם טוב',
    'גוני ברית בן': 'גוני בן ברית',
    '2ו איילת מלאך בוכניק': 'איילת מלאך בוכניק',
    '4\\2ה טלי נוראל': 'טלי נוראל'
}


def clean_cell_text(text):
    if not text:
        return ""
    clean = text.strip()
    clean = re.sub(r'[\r\n]+', '\n', clean)
    return clean


def set_cell_rtl(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(r'<w:tcMar %s><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>' % nsdecls('w'))
    tcPr.append(tcMar)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(r'<w:bidi %s/>' % nsdecls('w')))
        for r in p.runs:
            rPr = r._r.get_or_add_rPr()
            rPr.append(parse_xml(r'<w:rtl %s/>' % nsdecls('w')))


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def match_db_teacher(raw_name):
    # Direct special map
    if raw_name in SPECIAL_TEACHER_TITLES:
        return SPECIAL_TEACHER_TITLES[raw_name]
        
    cleaned = re.sub(r'^(?:חנמ\s*|מ[\"״׳\']?חנ\s*)?[1-4][א-ו](?:\\?[1-4][א-ו])?\s*', '', raw_name).strip()
    
    # 1. Exact match
    for dt in DB_TEACHERS:
        if cleaned == dt or cleaned.replace(" ", "") == dt.replace(" ", ""):
            return dt
            
    # 2. Unordered word match (e.g. "בן ברית גוני" <-> "גוני בן ברית", "סגל חן" <-> "חן סגל")
    clean_words = sorted(cleaned.split())
    for dt in DB_TEACHERS:
        dt_words = sorted(dt.split())
        if clean_words == dt_words:
            return dt
            
    # 3. Partial match
    for dt in DB_TEACHERS:
        if cleaned in dt or dt in cleaned:
            return dt
            
    return cleaned


def format_teacher_cell_content(cell_text):
    lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
    if not lines:
        return ""
    
    joined = ' '.join(lines)
    if 'פרטני' in joined:
        return "פרטני, פרטני"
    
    if 'שהייה' in joined:
        team_name = lines[0].replace('צוות', '').strip()
        if team_name == 'מורים מ"חנ':
            team_name = 'חנ"מ'
        elif team_name == 'א"מתי':
            team_name = 'מתי"א'
        return f"צוות {team_name}, שהייה" if team_name and team_name != 'שהייה' else "שהייה, שהייה"
    
    subject = lines[0] if len(lines) > 0 else ""
    if len(lines) == 2 and lines[1] in ['הוראה', 'שהייה']:
        class_part = ""
        activity = lines[1]
    elif len(lines) >= 3:
        class_part = lines[1]
        activity = lines[2]
    else:
        class_part = lines[1] if len(lines) > 1 else ""
        activity = "הוראה"
    
    # Fix subject corrections
    for k, v in PHRASE_CORRECTIONS.items():
        if subject == k:
            subject = v
            break
            
    # Format class_part
    if ',' in class_part:
        sub_classes = class_part.split(',')
        formatted_classes = []
        for sc in sub_classes:
            m = re.search(r'([1-4])([א-ו])', sc)
            if m:
                code_raw = f"{m.group(1)}{m.group(2)}"
                code_fmt = f"{m.group(2)}{m.group(1)}"
                hr = CLASS_HOMEROOM.get(code_raw, "")
                formatted_classes.append(f"{code_fmt} {hr}".strip())
            else:
                formatted_classes.append(sc.strip())
        final_class = ','.join(formatted_classes)
    elif class_part:
        m = re.search(r'([1-4])([א-ו])', class_part)
        if m:
            code_raw = f"{m.group(1)}{m.group(2)}"
            code_fmt = f"{m.group(2)}{m.group(1)}"
            hr = CLASS_HOMEROOM.get(code_raw, "")
            final_class = f"{code_fmt} {hr}".strip()
        else:
            final_class = class_part
    else:
        final_class = ""
            
    if final_class:
        return f"{subject}, {final_class}, {activity}"
    else:
        return f"{subject}, {activity}"


def format_class_cell_content(cell_text):
    lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
    if not lines:
        return ""
    
    joined = ' '.join(lines)
    if 'פרטני' in joined:
        return "פרטני, פרטני"
    if 'שהייה' in joined:
        return "שהייה, שהייה"
        
    subject = lines[0] if len(lines) > 0 else ""
    if len(lines) == 2 and lines[1] in ['הוראה', 'שהייה']:
        teacher = ""
        activity = lines[1]
    elif len(lines) >= 3:
        teacher = lines[1]
        activity = lines[2]
    else:
        teacher = lines[1] if len(lines) > 1 else ""
        activity = "הוראה"
    
    # Clean subject
    for k, v in PHRASE_CORRECTIONS.items():
        if subject == k:
            subject = v
            break
            
    # Match teacher with DB reference
    if teacher:
        teacher = match_db_teacher(teacher)
            
    if teacher:
        return f"{subject}, {teacher}, {activity}"
    else:
        return f"{subject}, {activity}"


# -------------------------------------------------------------
# 1. PROCESS CLASSES PDF -> כיתות.docx
# -------------------------------------------------------------
def process_classes():
    print(f"Opening Classes PDF: {CLASS_PDF}")
    doc_pdf = pymupdf.open(CLASS_PDF)
    doc_out = Document()
    
    for section in doc_out.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    class_pages = []
    for p_num in range(len(doc_pdf)):
        page = doc_pdf[p_num]
        tabs = page.find_tables()
        schedule_tabs = [t for t in tabs.tables if len(t.header.names) == 7]
        if not schedule_tabs:
            continue
        t = schedule_tabs[0]
        header_raw = page.get_text('text', clip=(0, 0, page.rect.width, t.bbox[1])).strip()
        
        m = re.search(r'([1-4])\s*([א-ו])', header_raw)
        if m:
            code_raw = f"{m.group(1)}{m.group(2)}"
            code_fmt = f"{m.group(2)}{m.group(1)}"
        else:
            code_raw = f"P{p_num+1}"
            code_fmt = f"כיתה {p_num+1}"
            
        hr_teacher = CLASS_HOMEROOM.get(code_raw, "")
        class_pages.append({
            'page_num': p_num,
            'code_raw': code_raw,
            'code_fmt': code_fmt,
            'hr_teacher': hr_teacher,
            'table': t,
            'page': page
        })

    grade_order = {'א': 1, 'ב': 2, 'ג': 3, 'ד': 4, 'ה': 5, 'ו': 6}
    def class_sort_key(cp):
        c = cp['code_fmt']
        if len(c) >= 2 and c[0] in grade_order:
            return (grade_order[c[0]], int(c[1]) if c[1].isdigit() else 99)
        return (99, 99)
        
    class_pages.sort(key=class_sort_key)
    print(f"Found {len(class_pages)} classes in PDF.")

    for cp in class_pages:
        page = cp['page']
        t = cp['table']
        code_fmt = cp['code_fmt']
        hr_teacher = cp['hr_teacher']
        
        # 1. Meta Header Table
        meta_table = doc_out.add_table(rows=1, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_cell_left = meta_table.cell(0, 0)
        meta_cell_right = meta_table.cell(0, 1)
        
        meta_cell_left.width = Inches(3.5)
        meta_cell_right.width = Inches(3.5)
        
        p_left = meta_cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_left = p_left.add_run(META_TIMESTAMP)
        r_left.font.size = Pt(8.5)
        r_left.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        
        p_right = meta_cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_right = p_right.add_run(SCHOOL_NAME)
        r_right.font.size = Pt(8.5)
        r_right.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        r_right.font.bold = True
        
        # 2. Title Paragraph
        p_title = doc_out.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p_title._p.get_or_add_pPr()
        pPr.append(parse_xml(r'<w:bidi %s/>' % nsdecls('w')))
        
        title_text = f"מערכת שעות כיתה {code_fmt} {hr_teacher}".strip()
        r_title = p_title.add_run(title_text)
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(14)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        # 3. Schedule Table
        sched_table = doc_out.add_table(rows=0, cols=2)
        sched_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(sched_table)
        
        col_to_day = {
            5: 1,  # ראשון
            4: 2,  # שני
            3: 3,  # שלישי
            2: 4,  # רביעי
            1: 5,  # חמישי
            0: 6   # שישי
        }
        
        day_lessons = {d_num: [] for d_num in range(1, 7)}
        
        for r_idx in range(1, len(t.rows)):
            row = t.rows[r_idx]
            hour_cell = row.cells[6] if len(row.cells) > 6 else None
            hour_text = page.get_text('text', clip=hour_cell).strip() if hour_cell else str(r_idx)
            m_h = re.search(r'(\d+)', hour_text)
            hour_num = int(m_h.group(1)) if m_h else r_idx
            
            for col_idx, day_num in col_to_day.items():
                if col_idx < len(row.cells) and row.cells[col_idx]:
                    raw_content = page.get_text('text', clip=row.cells[col_idx]).strip()
                    fmt_content = format_class_cell_content(raw_content)
                    if fmt_content:
                        day_lessons[day_num].append((hour_num, fmt_content))
                        
        for day_num, day_name in DAYS_ORDER:
            lessons = day_lessons[day_num]
            if not lessons:
                continue
                
            lessons.sort(key=lambda x: x[0])
            
            hdr_row = sched_table.add_row()
            c_left = hdr_row.cells[0]
            c_right = hdr_row.cells[1]
            
            c_left.width = Inches(5.2)
            c_right.width = Inches(2.0)
            
            set_cell_background(c_left, "28486B")
            set_cell_background(c_right, "28486B")
            
            p_rh = c_right.paragraphs[0]
            p_rh.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_rh = p_rh.add_run(day_name)
            r_rh.font.name = 'Calibri'
            r_rh.font.size = Pt(10.5)
            r_rh.font.bold = True
            r_rh.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
            set_cell_rtl(c_left)
            set_cell_rtl(c_right)
            
            for hour_num, content in lessons:
                l_row = sched_table.add_row()
                lc_left = l_row.cells[0]
                lc_right = l_row.cells[1]
                
                lc_left.width = Inches(5.2)
                lc_right.width = Inches(2.0)
                
                p_l = lc_left.paragraphs[0]
                p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_l = p_l.add_run(content)
                r_l.font.name = 'Calibri'
                r_l.font.size = Pt(9.5)
                r_l.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                
                p_r = lc_right.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                lbl = HOUR_LABELS.get(hour_num, f"שעה {hour_num}")
                r_r = p_r.add_run(lbl)
                r_r.font.name = 'Calibri'
                r_r.font.size = Pt(9.0)
                r_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                
                set_cell_rtl(lc_left)
                set_cell_rtl(lc_right)
                
        doc_out.add_paragraph()

    doc_out.save(CLASS_DOCX)
    print(f"Saved: {CLASS_DOCX}")


# -------------------------------------------------------------
# 2. PROCESS TEACHERS PDF -> מורים.docx
# -------------------------------------------------------------
def process_teachers():
    print(f"Opening Teachers PDF: {TEACHER_PDF}")
    doc_pdf = pymupdf.open(TEACHER_PDF)
    doc_out = Document()
    
    for section in doc_out.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    print(f"Found {len(doc_pdf)} teacher pages in PDF.")

    for p_num in range(len(doc_pdf)):
        page = doc_pdf[p_num]
        tabs = page.find_tables()
        schedule_tabs = [t for t in tabs.tables if len(t.header.names) == 7]
        if not schedule_tabs:
            continue
        t = schedule_tabs[0]
        header_raw = page.get_text('text', clip=(0, 0, page.rect.width, t.bbox[1])).strip()
        
        lines = [l.strip() for l in header_raw.split('\n') if l.strip()]
        title_line = ""
        for l in lines:
            if 'למורה שעות מערכת' in l or 'שעות מערכת' in l or 'מערכת שעות' in l:
                title_line = l
                break
        if not title_line and lines:
            title_line = lines[-1]
            
        raw_name = title_line.replace('למורה שעות מערכת', '').replace('מערכת שעות למורה', '').replace('מערכת שעות', '').strip()
        clean_name = match_db_teacher(raw_name)

        # 1. Meta Header Table
        meta_table = doc_out.add_table(rows=1, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_cell_left = meta_table.cell(0, 0)
        meta_cell_right = meta_table.cell(0, 1)
        
        meta_cell_left.width = Inches(3.5)
        meta_cell_right.width = Inches(3.5)
        
        p_left = meta_cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_left = p_left.add_run(META_TIMESTAMP)
        r_left.font.size = Pt(8.5)
        r_left.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        
        p_right = meta_cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_right = p_right.add_run(SCHOOL_NAME)
        r_right.font.size = Pt(8.5)
        r_right.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        r_right.font.bold = True
        
        # 2. Title Paragraph
        p_title = doc_out.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p_title._p.get_or_add_pPr()
        pPr.append(parse_xml(r'<w:bidi %s/>' % nsdecls('w')))
        
        title_text = f"מערכת שעות מורה {clean_name}".strip()
        r_title = p_title.add_run(title_text)
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(14)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        # 3. Schedule Table
        sched_table = doc_out.add_table(rows=0, cols=2)
        sched_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(sched_table)
        
        col_to_day = {
            5: 1,  # ראשון
            4: 2,  # שני
            3: 3,  # שלישי
            2: 4,  # רביעי
            1: 5,  # חמישי
            0: 6   # שישי
        }
        
        day_lessons = {d_num: [] for d_num in range(1, 7)}
        
        for r_idx in range(1, len(t.rows)):
            row = t.rows[r_idx]
            hour_cell = row.cells[6] if len(row.cells) > 6 else None
            hour_text = page.get_text('text', clip=hour_cell).strip() if hour_cell else str(r_idx)
            m_h = re.search(r'(\d+)', hour_text)
            hour_num = int(m_h.group(1)) if m_h else r_idx
            
            for col_idx, day_num in col_to_day.items():
                if col_idx < len(row.cells) and row.cells[col_idx]:
                    raw_content = page.get_text('text', clip=row.cells[col_idx]).strip()
                    fmt_content = format_teacher_cell_content(raw_content)
                    if fmt_content:
                        day_lessons[day_num].append((hour_num, fmt_content))
                        
        for day_num, day_name in DAYS_ORDER:
            lessons = day_lessons[day_num]
            if not lessons:
                continue
                
            lessons.sort(key=lambda x: x[0])
            
            hdr_row = sched_table.add_row()
            c_left = hdr_row.cells[0]
            c_right = hdr_row.cells[1]
            
            c_left.width = Inches(5.2)
            c_right.width = Inches(2.0)
            
            set_cell_background(c_left, "28486B")
            set_cell_background(c_right, "28486B")
            
            p_rh = c_right.paragraphs[0]
            p_rh.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_rh = p_rh.add_run(day_name)
            r_rh.font.name = 'Calibri'
            r_rh.font.size = Pt(10.5)
            r_rh.font.bold = True
            r_rh.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
            set_cell_rtl(c_left)
            set_cell_rtl(c_right)
            
            for hour_num, content in lessons:
                l_row = sched_table.add_row()
                lc_left = l_row.cells[0]
                lc_right = l_row.cells[1]
                
                lc_left.width = Inches(5.2)
                lc_right.width = Inches(2.0)
                
                p_l = lc_left.paragraphs[0]
                p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_l = p_l.add_run(content)
                r_l.font.name = 'Calibri'
                r_l.font.size = Pt(9.5)
                r_l.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                
                p_r = lc_right.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                lbl = HOUR_LABELS.get(hour_num, f"שעה {hour_num}")
                r_r = p_r.add_run(lbl)
                r_r.font.name = 'Calibri'
                r_r.font.size = Pt(9.0)
                r_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                
                set_cell_rtl(lc_left)
                set_cell_rtl(lc_right)
                
        doc_out.add_paragraph()

    doc_out.save(TEACHER_DOCX)
    print(f"Saved: {TEACHER_DOCX}")


if __name__ == "__main__":
    print("Starting conversion for Sharret Kiryat Ono (DB-aligned)...")
    process_classes()
    process_teachers()
    print("Done! Both כיתות.docx and מורים.docx have been successfully generated.")
