import os
import sys
import re
from datetime import datetime
import openpyxl

# התאמת קידוד פלט ל-UTF-8 בסביבת Windows
if sys.stdout:
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DAYS_ORDER = [
    ('ראשון', 'יום ראשון'),
    ('שני', 'יום שני'),
    ('שלישי', 'יום שלישי'),
    ('רביעי', 'יום רביעי'),
    ('חמישי', 'יום חמישי'),
    ('שישי', 'יום שישי'),
]

WORKGROUP_KEYWORDS = [
    "שילוב", "שהייה", "פרטני", "צוות", "ישיב", "ריכוז", "השתלמות", "ניהול", "תפקיד", "חלון", "הדרכה", "הכלה", "מליאה"
]

# מיפוי שמות פרטיים לשמות מלאים של ויצמן מתוך ה-DB
TEACHER_NAME_CANONICAL = {
    "אור": "אור סמרה",
    "אור סמרה": "אור סמרה",
    "אורטל": "אורטל באבו",
    "אורטל באבו": "אורטל באבו",
    "אורטל סבן": "אורטל סבן",
    "אורית": "אורית לנדאו",
    "אורית לנדאו": "אורית לנדאו",
    "אורנה": "אורנה מרקוביץ'",
    "אייל": "אייל חסידי",
    "אלה": "אלה רויטמן",
    "אנה": "אנה אשכנזי",
    "דיתי": "דיתי כהן",
    "דנית": "דנית אלבז",
    "הגר": "הגר עטר",
    "הדר": "הדר יצחק",
    "ורד": "ורד גולן",
    "זיו": "זיו ורבין",
    "חביבה": "חביבה רוזן",
    "טל": "טל רם",
    "טל רם": "טל רם",
    "טלי": "טלי טסלר",
    "טלי טסלר": "טלי טסלר",
    "טלי בצון": "טלי בצון",
    "יעל": "יעל בן סירה",
    "יעל בן סירה": "יעל בן סירה",
    "יעל חלבי": "יעל חלבי",
    "יפעת": "יפעת תומר",
    "יפעת תומר": "יפעת תומר",
    "יפעת יהודה": "יפעת יהודה",
    "כרמן": "כרמן ציטרון",
    "לאה": "לאה דיין",
    "מאי": "מאי מוליסיאן",
    "מאיה": "מאיה לוי",
    "מוריה": "מוריה גבאי",
    "מירי": "מירי טימסיט",
    "מלי": "מלי קרונקופ",
    "מעיין": "מעיין כהן",
    "מריאן": "מריאן זוהר",
    "מריאן זוהר": "מריאן זוהר",
    "נופר": "נופר מזרחי",
    "נחשון": "נחשון",
    "נטלי": "נטלי חג יחיא",
    "נעמה": "נעמה ביטון",
    "ספיר": "ספיר כהן אלשטיין",
    "ספיר כהן": "ספיר כהן אלשטיין",
    "סתיו": "סתיו גבאי",
    "עינב": "עינב ברי",
    "ענת": "ענת רימון",
    "פורל": "פורל רוט",
    "ציפי": "ציפי ארונפויד",
    "רויטל": "רויטל הראל",
    "רותי": "רותי זכריה",
    "רותי זכריה": "רותי זכריה",
    "רותם": "רותם אזולאי",
    "רחל": "רחל נמטלוב",
    "ריבה": "ריבה הראל",
    "ריזי": "ריזי יועצת",
    "שירה": "שירה ליברמן",
    "שירל": "שירל רוט",
    "שירלי": "שירלי ליאור",
    "שרון": "שרון דורנר גרינשטיין",
    "שריי": "שריי אוחנה",
}


def set_cell_margins(cell, top=5, bottom=5, left=50, right=50):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shd)


def set_rtl(paragraph_or_cell):
    pPr = paragraph_or_cell._element.get_or_add_pPr() if hasattr(paragraph_or_cell, '_element') else paragraph_or_cell
    if pPr.find(qn('w:bidi')) is None:
        pPr.append(OxmlElement('w:bidi'))


def clean_teacher_name_title(raw_title):
    clean = str(raw_title).strip()
    clean = re.sub(r'^מערכת שעות\s+(?:ל?מורה|מורה:?)\s*', '', clean)
    clean = re.sub(r'^מערכת שעות\s*', '', clean)
    
    # הסרת סיומת של כיתות כגון 'ו2+ו3', 'א3', 'מחנכת' וכו'
    clean = re.sub(r'\s+[א-ח][1-9](?:[\+,/][א-ח][1-9])*$', '', clean)
    clean = re.sub(r'\s+מחנכ(?:ת)?.*$', '', clean)
    clean = clean.strip()
    
    full_name = TEACHER_NAME_CANONICAL.get(clean, clean)
    return f"מערכת שעות מורה {full_name}"


def clean_class_name_title(raw_title):
    clean = str(raw_title).strip()
    clean = re.sub(r'^מערכת שעות\s+(?:ל?כיתה|כיתה:?)\s*', '', clean)
    clean = re.sub(r'^מערכת שעות\s*', '', clean)
    clean = clean.strip()
    if not clean.startswith("כיתה"):
        clean = f"כיתה {clean}"
    return f"מערכת שעות {clean}"


def format_hour_label(raw_hour):
    if not raw_hour:
        return ""
    text = str(raw_hour).strip()
    if "בוקר" in text or "08:00-08:15" in text:
        return "שעה 0 (08:00-08:15)"
    
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if len(lines) >= 2 and re.match(r'^\d+$', lines[0]) and '-' in lines[1]:
        return f"שעה {lines[0]} ({lines[1]})"
    
    match = re.match(r'^(\d+)\s*\((.*?)\)$', text)
    if match:
        return f"שעה {match.group(1)} ({match.group(2)})"
    
    if text.isdigit():
        return f"שעה {text}"
    
    return f"שעה {text}"


def format_class_lesson_cell(cell_value):
    if not cell_value:
        return ""
    text = str(cell_value).strip()
    text = re.sub(r'\bאמנות\b', 'אומנות', text)
    
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return ""
    
    if len(lines) == 1:
        val = lines[0]
        if any(kw in val for kw in WORKGROUP_KEYWORDS):
            return val
        
        # בדיקה אם יש שילוב מורים עם פלוס כגון 'אורית+ אורטל' או 'ספיר+ורד'
        if '+' in val:
            parts = [p.strip() for p in val.split('+')]
            if all(p in TEACHER_NAME_CANONICAL for p in parts):
                full_names = " / ".join(TEACHER_NAME_CANONICAL[p] for p in parts)
                return f"חינוך, {full_names}, הוראה"
        
        # אם יש ארוחת צהריים עם שם כגון 'נופר- ארוחת צהריים'
        if 'ארוחת צהריים' in val or 'ארוחת.צהריים' in val or 'ארוחת-צהריים' in val:
            for t_short, t_full in TEACHER_NAME_CANONICAL.items():
                if t_short in val:
                    return f"ארוחת צהריים, {t_full}, שהייה"
            return "ארוחת צהריים, שהייה"
        
        # אם יש 'אנגלית רחל' או 'ספיר אנגלית'
        if 'אנגלית' in val and any(t in val for t in TEACHER_NAME_CANONICAL):
            for t_short, t_full in TEACHER_NAME_CANONICAL.items():
                if t_short in val and t_short != "אנגלית":
                    return f"אנגלית, {t_full}, הוראה"
            return "אנגלית, הוראה"
        
        # אם זה רק שם מורה יחיד (כגון 'אורטל סבן', 'רותי', 'נעמה', 'אורית')
        if val in TEACHER_NAME_CANONICAL:
            full_teacher = TEACHER_NAME_CANONICAL[val]
            return f"חינוך, {full_teacher}, הוראה"
        
        return f"{val}, הוראה" if "הוראה" not in val else val
    
    # 2 שורות ומעלה (מקצוע + מורה)
    subject_part = lines[0]
    subject_part = re.sub(r'\bאמנות\b', 'אומנות', subject_part)
    teacher_part = lines[1]
    
    # הסרת סיומת כיתה משם המורה (כגון 'אלה רויטמן ג3' -> 'אלה רויטמן')
    teacher_part = re.sub(r'\s+[א-ח][1-9](?:[\+,/][א-ח][1-9])*$', '', teacher_part).strip()
    
    # ניקוי והשלמת שם מורה מלא
    if teacher_part in TEACHER_NAME_CANONICAL:
        teacher_part = TEACHER_NAME_CANONICAL[teacher_part]
    elif '+' in teacher_part:
        parts = [p.strip() for p in teacher_part.split('+')]
        if all(p in TEACHER_NAME_CANONICAL for p in parts):
            teacher_part = " / ".join(TEACHER_NAME_CANONICAL[p] for p in parts)
        
    joined = f"{subject_part}, {teacher_part}"
    if not any(kw in joined for kw in WORKGROUP_KEYWORDS) and "הוראה" not in joined:
        joined += ", הוראה"
    return joined


def format_teacher_lesson_cell(cell_value):
    if not cell_value:
        return ""
    text = str(cell_value).strip()
    text = re.sub(r'\bאמנות\b', 'אומנות', text)
    
    # תיקון שגיאות הקלדה ידועות
    text = text.replace('ג1י', 'ג1')
    text = text.replace('צוות שפה/,צוות סגני', 'צוות שפה / סגנים').replace('צוות שפה/', 'צוות שפה')
    
    # טיפול בהוראה מותאמת עם תלמיד וכיתה (כגון 'הוראה מותאמת- אור ליטבק ח2')
    match_horaa = re.match(r'הוראה מותאמת(?:-\s*.*?)?\s+([א-ח][1-9])$', text)
    if match_horaa:
        cls = match_horaa.group(1)
        return f"הוראה מותאמת, {cls}, הוראה"
    
    # טיפול ב'ישיבת צוות כיתה' כגון 'ישיבת צוות ב1'
    if re.match(r'ישיבת צוות\s+[א-ח][1-9]', text):
        return "ישיבת צוות, שהייה"
        
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return ""
    
    # אם זה תא של שורה 1
    if len(lines) == 1:
        val = lines[0]
        
        # קבוצות עבודה מפורשות
        if any(kw in val for kw in WORKGROUP_KEYWORDS):
            if "שהייה" not in val and "פרטני" not in val and "תפקיד" not in val and "השתלמות" not in val:
                return f"{val}, שהייה"
            return val
        
        # אם יש 'ב1 מדעים' או 'ב1 מולדת' או 'ג1 +ספיר'
        match_cls_sub = re.match(r'^([א-ח][1-9])\s*(.*)$', val)
        if match_cls_sub:
            cls = match_cls_sub.group(1)
            sub = match_cls_sub.group(2).strip()
            if not sub or sub.startswith('+') or any(t in sub for t in TEACHER_NAME_CANONICAL):
                return f"חינוך, {cls}, הוראה"
            return f"{sub}, {cls}, הוראה"
        
        # כיתה בלבד (כגון 'ח1')
        if re.match(r'^[א-ח][1-9]', val):
            cls_part = re.match(r'^([א-ח][1-9])', val).group(1)
            return f"חינוך, {cls_part}, הוראה"
        
        return f"{val}, הוראה" if "הוראה" not in val else val
    
    # 2 שורות ומעלה
    # בדיקה אם השורה הראשונה היא כיתה והשנייה מקצוע (הפוך! כגון 'ב1 \n מדעים')
    if re.match(r'^[א-ח][1-9]$', lines[0]) and not re.match(r'^[א-ח][1-9]$', lines[1]):
        cls_part = lines[0]
        sub_part = lines[1]
        return f"{sub_part}, {cls_part}, הוראה"
        
    # אם שורה שנייה היא שהייה/תפקיד/מליאה/פרטני
    if any(kw in lines[1] for kw in WORKGROUP_KEYWORDS):
        return f"{lines[0]}, {lines[1]}"
        
    # מקצוע + כיתה
    sub_part = lines[0]
    cls_part = lines[1]
    
    # ניקוי אם בשורה השנייה יש סיומת כגון 'הוראה'
    cls_match = re.search(r'([א-ח][1-9](?:\s*,\s*[א-ח][1-9])*)', cls_part)
    if cls_match:
        cls_part = cls_match.group(1)
        return f"{sub_part}, {cls_part}, הוראה"
        
    joined = f"{sub_part}, {cls_part}"
    if any(kw in joined for kw in WORKGROUP_KEYWORDS):
        return joined
    return f"{joined}, הוראה"


def parse_excel_schedules(excel_path, is_teacher_file=False):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active
    
    schedules = []
    r = 1
    
    while r <= ws.max_row:
        first_cell = ws.cell(row=r, column=1).value
        if first_cell and str(first_cell).strip().startswith("מערכת שעות"):
            raw_title = str(first_cell).strip()
            
            if is_teacher_file:
                title_text = clean_teacher_name_title(raw_title)
            else:
                title_text = clean_class_name_title(raw_title)
                
            r += 1
            header_row_vals = [ws.cell(row=r, column=c).value for c in range(1, 8)]
            day_col_map = {}
            for c_idx, val in enumerate(header_row_vals, start=1):
                if val:
                    val_str = str(val).strip()
                    for day_key, _ in DAYS_ORDER:
                        if day_key in val_str:
                            day_col_map[day_key] = c_idx
            
            r += 1
            day_schedule = {day_key: [] for day_key, _ in DAYS_ORDER}
            
            while r <= ws.max_row:
                hour_cell = ws.cell(row=r, column=1).value
                if hour_cell and str(hour_cell).strip().startswith("מערכת שעות"):
                    break
                
                row_vals = [ws.cell(row=r, column=c).value for c in range(1, 8)]
                if not any(v is not None for v in row_vals):
                    next_first = ws.cell(row=r+1, column=1).value if r+1 <= ws.max_row else None
                    if next_first and str(next_first).strip().startswith("מערכת שעות"):
                        r += 1
                        break
                    r += 1
                    continue
                
                if hour_cell:
                    hour_label = format_hour_label(hour_cell)
                    for day_key, _ in DAYS_ORDER:
                        if day_key in day_col_map:
                            col_idx = day_col_map[day_key]
                            cell_v = ws.cell(row=r, column=col_idx).value
                            if cell_v and str(cell_v).strip():
                                if is_teacher_file:
                                    formatted_text = format_teacher_lesson_cell(cell_v)
                                else:
                                    formatted_text = format_class_lesson_cell(cell_v)
                                
                                if formatted_text:
                                    day_schedule[day_key].append((hour_label, formatted_text))
                r += 1
            
            schedules.append({
                'title_text': title_text,
                'schedule': day_schedule
            })
        else:
            r += 1
            
    return schedules


def create_docx_from_schedules(output_docx_path, schedules, school_name="ויצמן"):
    out_doc = Document()

    for section in out_doc.sections:
        section.top_margin = Pt(11)
        section.bottom_margin = Pt(22)
        section.left_margin = Pt(22)
        section.right_margin = Pt(22)

    now_str = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    for idx, item in enumerate(schedules):
        meta_table = out_doc.add_table(rows=1, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        c0 = meta_table.rows[0].cells[0]
        c0.text = now_str
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.runs[0].font.name = 'Arial'
        p0.runs[0].font.size = Pt(10)
        
        c1 = meta_table.rows[0].cells[1]
        c1.text = school_name
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.runs[0].font.name = 'Arial'
        p1.runs[0].font.size = Pt(10)

        title_p = out_doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = title_p.add_run(item['title_text'])
        t_run.font.name = 'Arial'
        t_run.font.size = Pt(14)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(0, 0, 0)
        set_rtl(title_p)

        day_schedule = item['schedule']
        total_rows = 0
        active_days = []
        for day_key, day_title in DAYS_ORDER:
            lessons = day_schedule.get(day_key, [])
            if lessons:
                active_days.append((day_title, lessons))
                total_rows += 1 + len(lessons)

        if total_rows > 0:
            list_table = out_doc.add_table(rows=total_rows, cols=2)
            list_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}><w:gridCol w:w="6000"/><w:gridCol w:w="2000"/></w:tblGrid>')
            list_table._element.insert(1, tblGrid)

            curr_row_idx = 0
            for day_title, lessons in active_days:
                day_row = list_table.rows[curr_row_idx]
                
                c_left = day_row.cells[0]
                set_cell_shading(c_left, '28486B')
                set_cell_margins(c_left)
                set_rtl(c_left.paragraphs[0])
                
                c_right = day_row.cells[1]
                set_cell_shading(c_right, '28486B')
                set_cell_margins(c_right)
                p_r = c_right.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r_day = p_r.add_run(day_title)
                r_day.font.name = 'Arial'
                r_day.font.size = Pt(9)
                r_day.font.bold = True
                r_day.font.color.rgb = RGBColor(255, 255, 255)
                set_rtl(p_r)
                
                curr_row_idx += 1

                for hour_label, lesson_content in lessons:
                    lesson_row = list_table.rows[curr_row_idx]
                    
                    c_lesson = lesson_row.cells[0]
                    set_cell_shading(c_lesson, 'FFFFFF')
                    set_cell_margins(c_lesson)
                    p_les = c_lesson.paragraphs[0]
                    p_les.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_les = p_les.add_run(lesson_content)
                    r_les.font.name = 'Arial'
                    r_les.font.size = Pt(9)
                    r_les.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_les)

                    c_hour = lesson_row.cells[1]
                    set_cell_shading(c_hour, 'FFFFFF')
                    set_cell_margins(c_hour)
                    p_hr = c_hour.paragraphs[0]
                    p_hr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_hr = p_hr.add_run(hour_label)
                    r_hr.font.name = 'Arial'
                    r_hr.font.size = Pt(9)
                    r_hr.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_hr)

                    curr_row_idx += 1

        if idx < len(schedules) - 1:
            p_break = out_doc.add_paragraph()
            p_break.add_run().add_break(WD_BREAK.PAGE)

    out_doc.save(output_docx_path)
    print(f"נוצר בהצלחה: {output_docx_path}")


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    classes_excel = os.path.join(base_dir, "כיתות.xlsx")
    classes_docx = os.path.join(base_dir, "כיתות.docx")
    if os.path.exists(classes_excel):
        print(f"ממיר קובץ כיתות: {classes_excel}...")
        classes_schedules = parse_excel_schedules(classes_excel, is_teacher_file=False)
        print(f"נמצאו {len(classes_schedules)} מערכות שעות של כיתות.")
        create_docx_from_schedules(classes_docx, classes_schedules, school_name="ויצמן")
    
    teachers_excel = os.path.join(base_dir, "מורים.xlsx")
    teachers_docx = os.path.join(base_dir, "מורים.docx")
    if os.path.exists(teachers_excel):
        print(f"\nממיר קובץ מורים: {teachers_excel}...")
        teachers_schedules = parse_excel_schedules(teachers_excel, is_teacher_file=True)
        print(f"נמצאו {len(teachers_schedules)} מערכות שעות של מורים.")
        create_docx_from_schedules(teachers_docx, teachers_schedules, school_name="ויצמן")


if __name__ == "__main__":
    main()
