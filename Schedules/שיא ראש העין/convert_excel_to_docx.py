import os
import sys
import re
from datetime import datetime
import xlrd

# Adjust stdout encoding for Windows UTF-8
if sys.stdout:
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DAYS_MAP = {
    1: ('ראשון', 'יום ראשון'),
    2: ('שני', 'יום שני'),
    3: ('שלישי', 'יום שלישי'),
    4: ('רביעי', 'יום רביעי'),
    5: ('חמישי', 'יום חמישי'),
    6: ('שישי', 'יום שישי'),
}

HOUR_LABELS = {
    1: 'שעה 1 (08:00-08:50)',
    2: 'שעה 2 (08:50-09:35)',
    3: 'שעה 3 (10:15-11:00)',
    4: 'שעה 4 (11:00-11:45)',
    5: 'שעה 5 (12:00-12:45)',
    6: 'שעה 6 (12:45-13:30)',
    7: 'שעה 7 (13:30-14:20)',
    8: 'שעה 8',
    9: 'שעה 9',
}

SUBJECT_MAP = {
    'אוריי.מחשב': 'אוריינות מחשב',
    'אוריינות מ': 'אוריינות מחשב',
    'בשביל.המורש': 'בשבילי המורשת',
    'בשבילי המו': 'בשבילי המורשת',
    'גינה.לימודי': 'גינה לימודית',
    'גינה לימוד': 'גינה לימודית',
    'חינוך.גופני': 'חינוך גופני',
    'חינוך גופנ': 'חינוך גופני',
    'חינוך.פיננס': 'חינוך פיננסי',
    'חינוך פיננ': 'חינוך פיננסי',
    'כישור.חיים': 'כישורי חיים',
    'כישורי חיי': 'כישורי חיים',
    'משחקי.בעברי': 'משחקים בעברית',
    'משחקים בעב': 'משחקים בעברית',
    'עפים.על.העול': 'עפים על העולם',
    'עפים על הע': 'עפים על העולם',
    'תרבו.יהו.ישר': 'תרבות יהודית ישראלית',
    'תרבות יהוד': 'תרבות יהודית ישראלית',
}

WORKGROUP_KEYWORDS = [
    'פרטני', 'שהייה', 'רוחב', 'ניהול', 'ייעוץ', 'מליאה',
    'ישיבת צוות', 'ישיבת הנהלה מורחבת', 'צוות ניהול', 'הדרכה שפה'
]

TEACHER_LASTNAME_MAP = {
    '- בנצי יטי': 'קרב - בנצי יטיב',
    'אבו סנינה': 'שירין אבו סנינה שירין',
    'אברהם': 'שירה אברהם',
    'אלוני': 'עמית אלוני',
    'ביטון': 'מיכל ביטון',
    'גטה': 'תקווה גטה',
    'גמליאל': 'סתיו גמליאל',
    'הרשקוביץ': 'אורטל הרשקוביץ',
    'וויט': 'סיגל וויט',
    'זכריה': 'עדן זכריה',
    'חומן וכטל': 'מור חומן וכטל',
    'חכם': 'חן חכם',
    'חסון': 'שלי חסון',
    'ידאעי': 'שני ידאעי',
    'ימין': 'ליאונור ימין',
    'ינאי': 'עדי ינאי',
    'יריב': 'מיטל יריב',
    'ישועה': 'סימה ישועה',
    'לרנר': 'תמי לרנר',
    'מידז׳נסקי': 'עינב מידז׳נסקי איפרגן',
    'מסראוה': 'מרח מסראוה',
    'נדוף': 'אור נדוף',
    'ניסני': 'מזי ניסני',
    'סיני': 'שרון סיני',
    'ספיר': 'גינה-דרוד ספיר',
    'עאזם': 'מחמוד עאזם',
    'עוזרי': 'אור עוזרי',
    'פולק חנן': 'אלדמע פולק חנן',
    'פוקשנסקי פ': 'ליז פוקשנסקי פיינשטיין',
    'קליין': 'תמר קליין',
    'קרלין': 'דנה קרלין',
    'קשטן': 'לידר קשטן',
    'רבינוביץ': 'מירי רבינוביץ',
    'שונשיין': 'שירן שונשיין',
    'שחר': 'לירז שחר',
    'שחר אלבז': 'עוז שחר אלבז',
    'שישליאניקו': 'ליליה שישליאניקוב',
    'שרגורודסקי': 'לירון שרגורודסקי',
    'שרף': 'בר שרף',
}

CLASS_HOMEROOM_TEACHERS = {
    'א1': 'מור חומן וכטל',
    'א2': 'לירז שחר',
    'א3': 'שירן שונשיין',
    'א4': 'אור נדוף',
}


def normalize_subject(sub: str) -> str:
    sub = sub.strip().replace('\u200b', '')
    return SUBJECT_MAP.get(sub, sub)


def normalize_teacher_workgroup(val: str) -> str:
    clean = val.strip()
    clean = clean.replace('ישיב.הנה.מור (ש)', 'ישיבת הנהלה מורחבת, שהייה')
    clean = clean.replace('ישיבת.צוות (ש)', 'ישיבת צוות, שהייה')
    clean = clean.replace('צוות.ניהול (ש)', 'צוות ניהול, שהייה')
    clean = clean.replace('רוחב (ש)', 'רוחב, שהייה')
    clean = clean.replace('מליאה (ש)', 'מליאה, שהייה')
    clean = clean.replace('הדרכה שפה (ש)', 'הדרכה שפה, שהייה')
    clean = clean.replace('הדרכה שפה', 'הדרכה שפה, שהייה')
    if clean in ['פרטני', 'שהייה', 'רוחב', 'ניהול', 'ייעוץ', 'מליאה']:
        return f"{clean}, שהייה"
    if any(kw in clean for kw in WORKGROUP_KEYWORDS) and 'שהייה' not in clean:
        return f"{clean}, שהייה"
    return clean


def resolve_teacher_name(t_label: str, class_name: str, subject: str, full_teachers: list) -> str:
    t_label = t_label.strip()
    if not t_label:
        return ''
    if t_label == 'ברק':
        if class_name.startswith('א'):
            return 'ורד ברק'
        return 'דנה ברק'
    if t_label == 'פרץ':
        if 'מוסיקה' in subject:
            return 'מימי פרץ'
        return 'ספיר פרץ'
    if t_label in TEACHER_LASTNAME_MAP:
        return TEACHER_LASTNAME_MAP[t_label]
    if t_label in full_teachers:
        return t_label
    for ft in full_teachers:
        if t_label in ft or ft.endswith(t_label):
            return ft
    return t_label


def set_cell_margins(cell, top=5, bottom=5, left=50, right=50):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
    pPr.append(pBdr)


def parse_teachers_xls(file_path: str):
    wb = xlrd.open_workbook(file_path, formatting_info=False)
    sheet = wb.sheet_by_index(0)

    schedules = []
    curr_teacher = None
    day_schedule = {}

    for r in range(sheet.nrows):
        val0 = str(sheet.cell_value(r, 0)).strip()
        if val0.startswith('מערכת שעות למורה '):
            if curr_teacher and day_schedule:
                schedules.append({
                    'title': f"מערכת שעות מורה {curr_teacher}",
                    'schedule': day_schedule
                })
            curr_teacher = val0.replace('מערכת שעות למורה ', '').strip()
            day_schedule = {d: [] for d in range(1, 7)}
            continue

        if not curr_teacher or val0 == 'שעה':
            continue

        try:
            hour = int(float(val0))
        except (ValueError, TypeError):
            continue

        hour_label = HOUR_LABELS.get(hour, f'שעה {hour}')

        for c in range(1, 7):
            cell_v = str(sheet.cell_value(r, c)).strip()
            if not cell_v:
                continue

            lines = [l.strip() for l in cell_v.split('\n') if l.strip()]
            if not lines:
                continue

            line0 = lines[0]
            # Check if workgroup / stay hour
            wg_norm = normalize_teacher_workgroup(line0)
            if 'שהייה' in wg_norm or any(kw in wg_norm for kw in WORKGROUP_KEYWORDS):
                day_schedule[c].append((hour_label, wg_norm))
            elif len(lines) >= 2:
                sub = normalize_subject(line0)
                cls = lines[1].strip()
                lesson_text = f"{sub}, {cls}, הוראה"
                day_schedule[c].append((hour_label, lesson_text))
            else:
                # Single line
                norm_sub = normalize_subject(line0)
                lesson_text = f"{norm_sub}, הוראה"
                day_schedule[c].append((hour_label, lesson_text))

    if curr_teacher and day_schedule:
        schedules.append({
            'title': f"מערכת שעות מורה {curr_teacher}",
            'schedule': day_schedule
        })

    return schedules


def parse_classes_xls(file_path: str, full_teachers: list):
    wb = xlrd.open_workbook(file_path, formatting_info=False)
    sheet = wb.sheet_by_index(0)

    schedules = []
    curr_class = None
    day_schedule = {}

    for r in range(sheet.nrows):
        val0 = str(sheet.cell_value(r, 0)).strip()
        if val0.startswith('מערכת שעות לכיתה'):
            if curr_class and day_schedule:
                schedules.append({
                    'title': f"מערכת שעות כיתה {curr_class}",
                    'schedule': day_schedule
                })
            curr_class = val0.replace('מערכת שעות לכיתה', '').strip()
            day_schedule = {d: [] for d in range(1, 7)}
            continue

        if not curr_class or val0 == 'שעה':
            continue

        try:
            hour = int(float(val0))
        except (ValueError, TypeError):
            continue

        hour_label = HOUR_LABELS.get(hour, f'שעה {hour}')

        for c in range(1, 7):
            cell_v = str(sheet.cell_value(r, c)).strip()
            if not cell_v:
                continue

            lines = [l.strip().replace('\u200b', '') for l in cell_v.split('\n') if l.strip()]
            if not lines:
                continue

            sub = normalize_subject(lines[0])
            t_label = lines[1] if len(lines) > 1 else ''

            # Special split handling for yoga in grade 1 on Tuesday (day 3)
            if c == 3 and curr_class in ['א1', 'א2', 'א3', 'א4'] and sub == 'יוגה':
                homeroom = CLASS_HOMEROOM_TEACHERS.get(curr_class, '')
                split_sub = normalize_subject(t_label) if t_label else 'שפה'
                lesson_text = f"יוגה, יוגה - שרון עזרן, הוראה\n{split_sub}, {homeroom}, הוראה"
            else:
                teacher_full = resolve_teacher_name(t_label, curr_class, sub, full_teachers)
                if teacher_full:
                    lesson_text = f"{sub}, {teacher_full}, הוראה"
                else:
                    lesson_text = f"{sub}, הוראה"

            day_schedule[c].append((hour_label, lesson_text))

    if curr_class and day_schedule:
        schedules.append({
            'title': f"מערכת שעות כיתה {curr_class}",
            'schedule': day_schedule
        })

    return schedules


def create_docx(output_path: str, schedules: list, school_name: str = "שיא ראש העין"):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(11)
        section.bottom_margin = Pt(22)
        section.left_margin = Pt(22)
        section.right_margin = Pt(22)

    now_str = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    for idx, item in enumerate(schedules):
        # Header meta table
        meta_table = doc.add_table(rows=1, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        c0 = meta_table.rows[0].cells[0]
        c0.text = now_str
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.runs[0].font.name = 'Arial'
        p0.runs[0].font.size = Pt(10)

        c1 = meta_table.rows[0].cells[1]
        c1.text = school_name
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.runs[0].font.name = 'Arial'
        p1.runs[0].font.size = Pt(10)

        # Title paragraph
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = title_p.add_run(item['title'])
        t_run.font.name = 'Arial'
        t_run.font.size = Pt(14)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(0, 0, 0)
        set_rtl(title_p)

        # Day schedules
        day_sched = item['schedule']
        active_days = []
        total_rows = 0
        for day_num in range(1, 7):
            lessons = day_sched.get(day_num, [])
            if lessons:
                _, day_title = DAYS_MAP[day_num]
                active_days.append((day_title, lessons))
                total_rows += 1 + len(lessons)

        if total_rows > 0:
            list_table = doc.add_table(rows=total_rows, cols=2)
            list_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}><w:gridCol w:w="6000"/><w:gridCol w:w="2000"/></w:tblGrid>')
            list_table._element.insert(1, tblGrid)

            curr_row = 0
            for day_title, lessons in active_days:
                day_row = list_table.rows[curr_row]

                c_left = day_row.cells[0]
                set_cell_shading(c_left, '28486B')
                set_cell_margins(c_left)
                set_rtl(c_left.paragraphs[0])

                c_right = day_row.cells[1]
                set_cell_shading(c_right, '28486B')
                set_cell_margins(c_right)
                p_r = c_right.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r_day = p_r.add_run(day_title)
                r_day.font.name = 'Arial'
                r_day.font.size = Pt(9)
                r_day.font.bold = True
                r_day.font.color.rgb = RGBColor(255, 255, 255)
                set_rtl(p_r)

                curr_row += 1

                for hour_label, lesson_content in lessons:
                    lesson_row = list_table.rows[curr_row]

                    c_lesson = lesson_row.cells[0]
                    set_cell_shading(c_lesson, 'FFFFFF')
                    set_cell_margins(c_lesson)
                    p_les = c_lesson.paragraphs[0]
                    p_les.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_les = p_les.add_run(lesson_content)
                    r_les.font.name = 'Arial'
                    r_les.font.size = Pt(9)
                    r_les.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_les)

                    c_hour = lesson_row.cells[1]
                    set_cell_shading(c_hour, 'FFFFFF')
                    set_cell_margins(c_hour)
                    p_hr = c_hour.paragraphs[0]
                    p_hr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r_hr = p_hr.add_run(hour_label)
                    r_hr.font.name = 'Arial'
                    r_hr.font.size = Pt(9)
                    r_hr.font.color.rgb = RGBColor(0, 0, 0)
                    set_rtl(p_hr)

                    curr_row += 1

        if idx < len(schedules) - 1:
            p_break = doc.add_paragraph()
            p_break.add_run().add_break(WD_BREAK.PAGE)

    doc.save(output_path)
    print(f"נוצר בהצלחה: {output_path} ({len(schedules)} מערכות)")


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    teachers_xls = os.path.join(base_dir, "מורים.xls")
    classes_xls = os.path.join(base_dir, "כיתות.xls")

    teachers_docx = os.path.join(base_dir, "מורים.docx")
    classes_docx = os.path.join(base_dir, "כיתות.docx")

    print(f"טוען קובץ מורים: {teachers_xls}...")
    teachers_schedules = parse_teachers_xls(teachers_xls)
    print(f"נמצאו {len(teachers_schedules)} מערכות שעות של מורים.")
    create_docx(teachers_docx, teachers_schedules, school_name="שיא ראש העין")

    full_teacher_names = [s['title'].replace('מערכת שעות מורה ', '').strip() for s in teachers_schedules]

    print(f"\nטוען קובץ כיתות: {classes_xls}...")
    classes_schedules = parse_classes_xls(classes_xls, full_teacher_names)
    print(f"נמצאו {len(classes_schedules)} מערכות שעות של כיתות.")
    create_docx(classes_docx, classes_schedules, school_name="שיא ראש העין")


if __name__ == "__main__":
    main()
